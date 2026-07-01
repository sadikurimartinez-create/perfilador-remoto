import docx
import sys

# Reconfigure stdout to use utf-8
sys.stdout.reconfigure(encoding='utf-8')

def read_docx(file_path):
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
    return '\n'.join(fullText)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_doc.py <file_path>")
        sys.exit(1)
    
    path = sys.argv[1]
    text = read_docx(path)
    print(text[:5000])  # Print first 5000 chars
